import sys
import os
import io
from docx import Document
from app.services.word_service import extract_schema_from_docx

def create_mock_docx() -> bytes:
    doc = Document()
    
    # 1. Tabla de Logotipo / Encabezado (Tabla 1)
    tbl1 = doc.add_table(rows=1, cols=3)
    cells = tbl1.rows[0].cells
    cells[0].text = "LOGO"
    cells[1].text = "ACTA DE REUNIÓN"
    cells[2].text = "Código: SGI-12\nVersión: 03"
    
    # 2. Tabla de Metadatos (Tabla 2)
    tbl2 = doc.add_table(rows=4, cols=4)
    tbl2_data = [
        ["Tipo:", "", "Comité:", ""],
        ["Fecha:", "", "Área:", ""],
        ["Hora Inicio:", "", "Hora Final:", ""],
        ["Lugar:", "", "Nro. Acta:", ""]
    ]
    for r_idx, row in enumerate(tbl2.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = tbl2_data[r_idx][c_idx]
            
    # 3. Encabezados de secciones
    p1 = doc.add_paragraph()
    p1.style = "Heading 1"
    p1.text = "1. PARTICIPANTES"
    
    # Tabla de participantes
    tbl3 = doc.add_table(rows=2, cols=4)
    tbl3.rows[0].cells[0].text = "ENTIDAD"
    tbl3.rows[0].cells[1].text = "NOMBRE"
    tbl3.rows[0].cells[2].text = "CARGO"
    tbl3.rows[0].cells[3].text = "FIRMA"
    
    tbl3.rows[1].cells[0].text = ""
    tbl3.rows[1].cells[1].text = ""
    tbl3.rows[1].cells[2].text = ""
    tbl3.rows[1].cells[3].text = ""
    
    p2 = doc.add_paragraph()
    p2.style = "Heading 1"
    p2.text = "2. DESARROLLO"
    
    doc.add_paragraph("Texto de desarrollo aquí...")
    
    p3 = doc.add_paragraph()
    p3.style = "Heading 1"
    p3.text = "3. PRÓXIMA REUNIÓN"
    
    doc.add_paragraph("Texto de próxima reunión...")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def run_tests():
    file_bytes = create_mock_docx()
    schema = extract_schema_from_docx(file_bytes)
    
    print("Schema extraído:")
    sections = schema.get("sections", [])
    for s in sections:
        print(f"- Sección: {s['title']} (clave: {s['key']})")
        for f in s.get("fields", []):
            print(f"  * Campo: {f['key']} | label: {f['label']} | tipo: {f['type']}")
            
    # Verificaciones
    assert len(sections) > 0, "No se extrajeron secciones"
    
    # 1. Verificar que la sección encabezado existe y tiene las claves correctas
    enc_sec = next((s for s in sections if s["key"] == "encabezado"), None)
    assert enc_sec is not None, "Falta la sección de encabezado"
    
    expected_fields = {"tipo", "comité", "fecha", "área", "hora_inicio", "hora_final", "lugar", "nro_acta"}
    extracted_fields = {f["key"] for f in enc_sec["fields"]}
    
    print(f"\nCampos esperados: {expected_fields}")
    print(f"Campos extraídos: {extracted_fields}")
    
    missing_fields = expected_fields - extracted_fields
    assert not missing_fields, f"Faltan campos de metadatos: {missing_fields}"
    
    # 2. Verificar que no se coló basura del logotipo (como código o versión)
    junk_fields = {"logo", "acta_de_reunion", "codigo", "version"} & extracted_fields
    assert not junk_fields, f"Se extrajeron campos basura del logotipo: {junk_fields}"
    
    # 3. Verificar que la sección "PRÓXIMA REUNIÓN" no fue omitida
    proxima_sec = next((s for s in sections if "próx" in s["key"] or "prox" in s["key"]), None)
    assert proxima_sec is not None, "Se omitió la sección de Próxima Reunión"
    
    # 4. Verificar el funcionamiento del nuevo parseador de strings a columnas
    from app.services.word_service import _parse_string_item_to_cols
    
    # 4.1. Participantes con comas (Nombre, Cargo, Entidad en UI)
    # Nota: la UI mapea Nombre, Cargo, Entidad, por lo que el split de comas da partes en ese orden.
    # En la plantilla, col_keys son ["entidad", "nombre", "cargo"]. 
    # Si la UI guardó el dict correcto, se rellenará vía isinstance(item, dict) directamente.
    # Pero si se guardó como string plano separado por comas, se procesa secuencialmente:
    res = _parse_string_item_to_cols("Innovasoft, Juan Pérez, Desarrollador", ["entidad", "nombre", "cargo"])
    assert res["entidad"] == "Innovasoft", f"Entidad incorrecta: {res['entidad']}"
    assert res["nombre"] == "Juan Pérez", f"Nombre incorrecto: {res['nombre']}"
    assert res["cargo"] == "Desarrollador", f"Cargo incorrecto: {res['cargo']}"

    # 4.2. Participante único sin comas (debe ir a nombre, no a entidad)
    res = _parse_string_item_to_cols("Juan Pérez", ["entidad", "nombre", "cargo"])
    assert res["nombre"] == "Juan Pérez", f"Nombre incorrecto: {res['nombre']}"
    assert res["entidad"] == "", f"Entidad no debería estar llena: {res['entidad']}"

    # 4.3. Compromisos con comas (Actividad, Responsable, Fecha)
    res = _parse_string_item_to_cols("Documentar el proceso, Pedro Gómez, 10 de mayo", ["actividad", "estado", "responsable", "fecha"])
    assert res["actividad"] == "Documentar el proceso", f"Actividad incorrecta: {res['actividad']}"
    assert res["responsable"] == "Pedro Gómez", f"Responsable incorrecto: {res['responsable']}"
    assert res["fecha"] == "10 de mayo", f"Fecha incorrecta: {res['fecha']}"

    # 4.4. Compromiso con dos puntos (Pedro Gómez: Documentar el proceso el 10 de mayo)
    res = _parse_string_item_to_cols("Pedro Gómez: Documentar el proceso para el 10 de mayo", ["actividad", "estado", "responsable", "fecha"])
    assert res["actividad"] == "Documentar el proceso", f"Actividad incorrecta: {res['actividad']}"
    assert res["responsable"] == "Pedro Gómez", f"Responsable incorrecto: {res['responsable']}"
    assert res["fecha"] == "10 de mayo", f"Fecha incorrecta: {res['fecha']}"

    # 4.5. Compromiso plano con extracción de nombre por patrón
    res = _parse_string_item_to_cols("Pedro Gómez se compromete a diseñar la interfaz el 15 de junio", ["actividad", "estado", "responsable", "fecha"])
    assert res["actividad"] == "Diseñar la interfaz", f"Actividad incorrecta: {res['actividad']}"
    assert res["responsable"] == "Pedro Gómez", f"Responsable incorrecto: {res['responsable']}"
    assert res["fecha"] == "15 de junio", f"Fecha incorrecta: {res['fecha']}"

    print("\n[OK] Todos los tests unitarios pasaron con exito!")

if __name__ == "__main__":
    run_tests()
