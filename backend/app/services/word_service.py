"""Extracción de schema y renderizado de plantillas Word.

Soporta dos tipos de plantilla:
  - Jinja2: contienen {{ variable }} — se extraen las variables y su contexto.
  - Estructural: plantillas corporativas con tablas y encabezados en MAYÚSCULAS,
    sin placeholders — se analiza la estructura (tablas + secciones) para
    generar un schema rico que la IA pueda llenar desde la transcripción.
"""
import io
import re
import copy
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn


# ─────────────────────────────────────────────────────────────────────────────
# Entrada pública
# ─────────────────────────────────────────────────────────────────────────────

def extract_schema_from_docx(file_bytes: bytes) -> dict:
    tpl = DocxTemplate(io.BytesIO(file_bytes))
    jinja_vars = tpl.get_undeclared_template_variables()
    if jinja_vars:
        schema = _extract_jinja2_schema(file_bytes, jinja_vars)
        schema["_tpl_type"] = "jinja2"
    else:
        schema = _extract_structural_schema(file_bytes)
        schema["_tpl_type"] = "structural"
    return schema


def render_acta_to_docx(template_path: str, acta_data: dict, tpl_type: str = "auto") -> bytes:
    """Elige el método de renderizado según el tipo de plantilla."""
    if tpl_type == "auto":
        with open(template_path, "rb") as f:
            raw = f.read()
        tpl = DocxTemplate(io.BytesIO(raw))
        tpl_type = "jinja2" if tpl.get_undeclared_template_variables() else "structural"

    if tpl_type == "jinja2":
        return _render_jinja2(template_path, acta_data)
    else:
        return _render_structural(template_path, acta_data)


# ─────────────────────────────────────────────────────────────────────────────
# Extracción: plantillas Jinja2
# ─────────────────────────────────────────────────────────────────────────────

def _extract_jinja2_schema(file_bytes: bytes, all_vars: set) -> dict:
    doc = Document(io.BytesIO(file_bytes))
    var_info = {v: {"surrounding": "", "section": "General"} for v in all_vars}
    section_order = ["General"]
    section_vars: dict[str, list] = {"General": []}
    current_section = "General"

    def _process(text: str):
        for var in re.findall(r"\{\{\s*(\w+)\s*\}\}", text):
            if var not in all_vars:
                continue
            clean = re.sub(r"\{\{[^}]*\}\}", "", text).strip(" :-–—\t\n")
            if not var_info[var]["surrounding"]:
                var_info[var]["surrounding"] = clean[:150]
            var_info[var]["section"] = current_section
            bucket = section_vars.setdefault(current_section, [])
            if var not in bucket:
                bucket.append(var)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if any(k in style for k in ("Heading", "Título", "Title", "Header")):
            current_section = text
            if current_section not in section_order:
                section_order.append(current_section)
            section_vars.setdefault(current_section, [])
        else:
            _process(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in table.cells:
                for para in cell.paragraphs:
                    _process(para.text.strip())

    assigned = {v for vs in section_vars.values() for v in vs}
    for v in sorted(all_vars):
        if v not in assigned:
            section_vars["General"].append(v)

    sections = []
    for sec_title in section_order:
        vars_here = section_vars.get(sec_title, [])
        if not vars_here:
            continue
        fields = []
        for var in vars_here:
            sur = var_info[var]["surrounding"]
            ftype = _infer_type(var, sur)
            label = sur if sur else var.replace("_", " ").title()
            fields.append({
                "key": var, "label": label[:100], "type": ftype,
                "required": _is_required(var, sur),
                "hint": _build_hint(var, sur, ftype),
            })
        sec_key = _to_key(sec_title)
        sections.append({"key": sec_key, "title": sec_title, "fields": fields})

    if not sections:
        sections = [{"key": "contenido", "title": "Contenido", "fields": [
            {"key": v, "label": v.replace("_", " ").title(),
             "type": _infer_type(v, ""), "required": _is_required(v, ""),
             "hint": _build_hint(v, "", _infer_type(v, ""))}
            for v in sorted(all_vars)
        ]}]

    return {"sections": sections}


# ─────────────────────────────────────────────────────────────────────────────
# Extracción: plantillas estructurales (tablas corporativas)
# ─────────────────────────────────────────────────────────────────────────────

def _is_valid_metadata_label(text: str) -> bool:
    clean = text.strip()
    if not clean or len(clean) < 2 or len(clean) > 40:
        return False
    lower_clean = clean.lower().rstrip(":")
    # Evitar campos estáticos de control de documentos
    if lower_clean in ("código", "codigo", "versión", "version", "pag", "pág", "página", "pagina", "proceso", "vigencia"):
        return False
    if clean.endswith(":"):
        return True
    keywords = {"fecha", "hora", "lugar", "comité", "comite", "tipo", "acta", "área", "area",
                "nombre", "cargo", "entidad", "participante", "asistente", "organiza", "convoca"}
    if any(k in lower_clean for k in keywords):
        return True
    return False


def _extract_structural_schema(file_bytes: bytes) -> dict:
    """Lee el documento elemento a elemento (párrafos y tablas en orden)
    para construir un schema fiel a la estructura corporativa."""
    doc = Document(io.BytesIO(file_bytes))
    body = doc.element.body
    para_tag, tbl_tag = qn("w:p"), qn("w:tbl")

    # Índices para correlacionar elementos XML con objetos python-docx
    para_list = list(doc.paragraphs)
    tbl_list = list(doc.tables)
    para_idx = tbl_idx = 0

    elements: list[tuple[str, object]] = []
    for child in body:
        if child.tag == para_tag and para_idx < len(para_list):
            elements.append(("para", para_list[para_idx]))
            para_idx += 1
        elif child.tag == tbl_tag and tbl_idx < len(tbl_list):
            elements.append(("table", tbl_list[tbl_idx]))
            tbl_idx += 1

    sections: list[dict] = []
    header_fields = []
    first_heading_seen = False
    i = 0

    while i < len(elements):
        kind, elem = elements[i]

        # ── Párrafo de encabezado ───────────────────────────────────────────
        if kind == "para" and _is_section_heading(elem):
            first_heading_seen = True
            heading = elem.text.strip().rstrip(".").strip()

            # ¿Le sigue una tabla?
            if i + 1 < len(elements) and elements[i + 1][0] == "table":
                sec = _parse_table_section(heading, elements[i + 1][1])
                if sec:
                    sections.append(sec)
                i += 2
            else:
                sec = _make_text_section(heading)
                if sec:
                    sections.append(sec)
                i += 1
            continue

        # ── Tablas de metadatos al inicio ──────────────────────────────────
        if kind == "table" and not first_heading_seen:
            fields = _parse_header_table(elem)
            if fields:
                header_fields.extend(fields)
            i += 1
            continue

        i += 1

    if header_fields:
        # Eliminar duplicados por clave manteniendo el orden
        seen_keys = set()
        unique_header_fields = []
        for f in header_fields:
            if f["key"] not in seen_keys:
                seen_keys.add(f["key"])
                unique_header_fields.append(f)
        sections.insert(0, {
            "key": "encabezado",
            "title": "Datos de la Reunión",
            "fields": unique_header_fields,
        })

    return {"sections": sections}


def _is_section_heading(para) -> bool:
    text = para.text.strip()
    if not text or len(text) < 3:
        return False
    style = para.style.name if para.style else ""
    if any(k in style for k in ("Heading", "Título", "Title")):
        return True
    clean = text.rstrip(".")
    return clean.isupper() and len(clean) > 3


def _parse_header_table(tbl) -> list:
    """Tabla tipo: | Etiqueta: | [valor] | Etiqueta: | [valor] |"""
    fields, seen = [], set()
    for row in tbl.rows:
        cells = row.cells
        j = 0
        while j < len(cells) - 1:
            label_text = cells[j].text.strip()
            if _is_valid_metadata_label(label_text):
                label = label_text.rstrip(":").strip()
                if label and label not in seen:
                    seen.add(label)
                    key = _to_key(label)
                    ftype = _infer_type(key, label)
                    fields.append({
                        "key": key,
                        "label": label,
                        "type": ftype,
                        "required": _is_required(key, label),
                        "hint": _build_hint(key, label, ftype),
                    })
            j += 2
    return fields


def _parse_table_section(heading: str, tbl) -> dict | None:
    """Convierte una sección con tabla en campos del schema según los headers.
    Genera _columns para TODAS las categorías de tabla para que el prompt
    de IA y la previsualización usen exactamente las columnas de la plantilla."""
    if not tbl.rows:
        return None
    headers = [c.text.strip() for c in tbl.rows[0].cells]
    h_join = " ".join(h.lower() for h in headers if h)
    sec_key = _to_key(heading)

    # ── Tabla de participantes ─────────────────────────────────────────────
    if any(k in h_join for k in ("nombre", "entidad", "firma", "participante", "cargo")):
        cols = [h for h in headers if h and "firma" not in h.lower()]
        col_keys = [_to_key(c) for c in cols]
        example_obj = ", ".join(f'"{k}": "..."' for k in col_keys)
        return {
            "key": sec_key, "title": heading,
            "fields": [{
                "key": sec_key,
                "label": f"Participantes ({' | '.join(cols)})",
                "type": "people",
                "required": True,
                "hint": (
                    f"Extrae de la transcripción TODOS los participantes presentes. "
                    f"La tabla tiene las columnas: {', '.join(cols)}. "
                    f"Para CADA participante, devuelve un objeto con EXACTAMENTE "
                    f"estas claves: [{{{example_obj}}}]. "
                    f"Si no encuentras un dato para alguna columna, pon \"—\"."
                ),
                "_table_type": "participants",
                "_columns": cols,
            }],
        }

    # ── Tabla de compromisos / actividades ────────────────────────────────
    if any(k in h_join for k in ("actividad", "responsable", "estado", "tarea", "comprom")):
        cols = [h for h in headers if h]
        col_keys = [_to_key(c) for c in cols]
        example_obj = ", ".join(f'"{k}": "..."' for k in col_keys)
        return {
            "key": sec_key, "title": heading,
            "fields": [{
                "key": sec_key,
                "label": f"Compromisos ({' | '.join(cols)})",
                "type": "list",
                "required": False,
                "hint": (
                    f"Extrae los compromisos o tareas de la transcripción. "
                    f"La tabla tiene las columnas: {', '.join(cols)}. "
                    f"Para CADA compromiso, devuelve un objeto con EXACTAMENTE "
                    f"estas claves: [{{{example_obj}}}]. "
                    f"Si no encuentras un dato para alguna columna, pon \"—\"."
                ),
                "_table_type": "commitments",
                "_columns": cols,
            }],
        }

    # ── Tabla de entregables ──────────────────────────────────────────────
    if (any(k in h_join for k in ("entregable", "material")) or
            ("tipo" in h_join and "tema" not in h_join)):
        cols = [h for h in headers if h]
        col_keys = [_to_key(c) for c in cols]
        example_obj = ", ".join(f'"{k}": "..."' for k in col_keys)
        return {
            "key": sec_key, "title": heading,
            "fields": [{
                "key": sec_key,
                "label": f"Entregables ({' | '.join(cols)})",
                "type": "list",
                "required": False,
                "hint": (
                    f"Extrae los documentos, materiales o entregables mencionados. "
                    f"La tabla tiene las columnas: {', '.join(cols)}. "
                    f"Para CADA entregable, devuelve un objeto con EXACTAMENTE "
                    f"estas claves: [{{{example_obj}}}]. "
                    f"Si no encuentras un dato para alguna columna, pon \"—\"."
                ),
                "_table_type": "deliverables",
                "_columns": cols,
            }],
        }

    # ── Tabla de temas / agenda ───────────────────────────────────────────
    if any(k in h_join for k in ("tema", "descripci", "punto", "agenda")):
        cols = [h for h in headers if h]
        col_keys = [_to_key(c) for c in cols]
        example_obj = ", ".join(f'"{k}": "..."' for k in col_keys)
        return {
            "key": sec_key, "title": heading,
            "fields": [{
                "key": sec_key,
                "label": f"Temas ({' | '.join(cols)})",
                "type": "list",
                "required": False,
                "hint": (
                    f"Extrae los temas o puntos de agenda abordados. "
                    f"La tabla tiene las columnas: {', '.join(cols)}. "
                    f"Para CADA tema, devuelve un objeto con EXACTAMENTE "
                    f"estas claves: [{{{example_obj}}}]. "
                    f"Si no encuentras un dato para alguna columna, pon \"—\"."
                ),
                "_table_type": "topics",
                "_columns": cols,
            }],
        }

    # ── Tabla genérica (tratar como tabla de registros con _columns) ──────
    cols = [h for h in headers if h]
    if not cols:
        return None
    col_keys = [_to_key(c) for c in cols]
    example_obj = ", ".join(f'"{k}": "..."' for k in col_keys)
    return {
        "key": sec_key, "title": heading,
        "fields": [{
            "key": sec_key,
            "label": f"{heading} ({' | '.join(cols)})",
            "type": "list",
            "required": False,
            "hint": (
                f"Extrae la información correspondiente a esta sección. "
                f"La tabla tiene las columnas: {', '.join(cols)}. "
                f"Para CADA registro, devuelve un objeto con EXACTAMENTE "
                f"estas claves: [{{{example_obj}}}]. "
                f"Si no encuentras un dato para alguna columna, pon \"—\"."
            ),
            "_table_type": "generic",
            "_columns": cols,
        }],
    }


def _make_text_section(heading: str) -> dict | None:
    h_low = heading.lower()
    skip = {"evidencias", "firma", "elaboro", "elaboró", "aprobado"}
    if any(k in h_low for k in skip):
        return None
    sec_key = _to_key(heading)
    ftype = "people" if any(k in h_low for k in ("participante", "asistente")) else \
            "list" if any(k in h_low for k in ("agenda", "puntos", "temas", "anexo")) else "text"
    return {
        "key": sec_key, "title": heading,
        "fields": [{
            "key": sec_key,
            "label": heading.title(),
            "type": ftype,
            "required": any(k in h_low for k in ("objetivo", "desarrollo", "proposit", "alcance")),
            "hint": _build_hint(sec_key, heading, ftype),
        }],
    }


# ─────────────────────────────────────────────────────────────────────────────
# Renderizado
# ─────────────────────────────────────────────────────────────────────────────

def _render_jinja2(template_path: str, acta_data: dict) -> bytes:
    doc = DocxTemplate(template_path)
    context: dict = {}
    for sec_val in acta_data.values():
        for field_key, field_val in sec_val.get("fields", {}).items():
            value = field_val.get("value", "") or ""
            if isinstance(value, list):
                value = ", ".join(str(v) for v in value)
            context[field_key] = str(value)
    doc.render(context)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _render_structural(template_path: str, acta_data: dict) -> bytes:
    """Rellena plantillas corporativas usando python-docx.
    Usa un índice por tipo de tabla para tolerar múltiples secciones del mismo tipo."""
    doc = Document(template_path)
    tables = doc.tables

    # Solo valores escalares (text/date) en flat — listas y personas van solo a typed
    # para evitar que _fill_text_paragraphs las vuelque como texto crudo
    flat: dict = {}
    for sec_val in acta_data.values():
        for fk, fv in sec_val.get("fields", {}).items():
            if fk not in flat:
                ftype = fv.get("type", "text")
                if ftype not in ("people", "list"):
                    flat[fk] = fv.get("value")

    # Índice tipado: agrupa valores por categoría semántica en orden de sección
    typed: dict[str, list] = {"people": [], "commits": [], "topics": [], "deliverables": []}
    for sec_val in acta_data.values():
        for fk, fv in sec_val.get("fields", {}).items():
            v = fv.get("value")
            if not v:
                continue
            ftype = fv.get("type", "")
            flabel = (fv.get("label", "") + " " + fk).lower()
            if ftype == "people":
                typed["people"].append(v)
            elif ftype == "list":
                if any(k in flabel for k in ("comprom", "actividad", "tarea", "responsabl")):
                    typed["commits"].append(v)
                elif any(k in flabel for k in ("entregable", "document", "material")):
                    typed["deliverables"].append(v)
                else:
                    typed["topics"].append(v)

    # Punteros para ir consumiendo valores de cada categoría en orden
    ptr = {k: 0 for k in typed}

    # ── Rellenar tablas de metadata (buscando etiquetas en cualquier tabla) ──
    for tbl in tables:
        for row in tbl.rows:
            cells = row.cells
            j = 0
            while j < len(cells) - 1:
                label = cells[j].text.strip().rstrip(":").strip()
                key = _to_key(label)
                if key in flat:
                    value = flat[key]
                    if value is not None and not cells[j + 1].text.strip():
                        _set_cell_text(cells[j + 1], str(value))
                j += 2

    # ── Tablas: rellenar por tipo de tabla (ej. listas, participantes) ──────────────────
    for tbl in tables:
        if not tbl.rows:
            continue
        headers = [c.text.strip().lower() for c in tbl.rows[0].cells]
        h_join = " ".join(headers)

        # Evitar procesar la tabla de metadatos o logotipo como si fuera una tabla de lista
        if any(k in h_join for k in ("tipo:", "comité:", "comite:", "lugar:", "fecha:")):
            continue

        if any(k in h_join for k in ("nombre", "entidad", "participante")):
            vals = typed["people"]
            if ptr["people"] < len(vals):
                _fill_list_table(tbl, vals[ptr["people"]], ["entidad", "nombre", "cargo"])
                ptr["people"] += 1

        elif any(k in h_join for k in ("actividad", "responsable", "estado")):
            vals = typed["commits"]
            if ptr["commits"] < len(vals):
                _fill_list_table(tbl, vals[ptr["commits"]], ["actividad", "estado", "responsable", "fecha"])
                ptr["commits"] += 1

        elif any(k in h_join for k in ("entregable", "material")) or \
                ("tipo" in h_join and "tema" not in h_join):
            vals = typed["deliverables"]
            if ptr["deliverables"] < len(vals):
                _fill_list_table(tbl, vals[ptr["deliverables"]], ["descripcion", "tipo"])
                ptr["deliverables"] += 1

        elif any(k in h_join for k in ("tema",)):
            vals = typed["topics"]
            if ptr["topics"] < len(vals):
                _fill_list_table(tbl, vals[ptr["topics"]], ["tema", "descripcion"])
                ptr["topics"] += 1

    # ── Párrafos de texto libre (OBJETIVO, DESARROLLO, etc.) ─────────────
    _fill_text_paragraphs(doc, flat)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _parse_string_item_to_cols(item: str, col_keys: list[str]) -> dict:
    """Intenta estructurar un string en columnas según palabras clave y formato."""
    res = {}
    for k in col_keys:
        res[k] = ""

    clean_item = item.strip()
    if not clean_item:
        return res

    primary_key = col_keys[0]
    is_text_primary = any(x in primary_key for x in ("actividad", "descripci", "tema", "comprom", "tarea", "descripcion"))

    # Caso 1: Separado por comas o punto y coma
    separator = ";" if ";" in item else ","
    if separator in clean_item:
        parts = [p.strip() for p in clean_item.split(separator) if p.strip()]
        if is_text_primary:
            # ── Mapeo inteligente por clasificación ──
            assigned_parts = {}
            # 1. Buscar fecha
            date_key = next((k for k in col_keys if "fecha" in k or "plazo" in k or "limite" in k), None)
            if date_key:
                for part in list(parts):
                    if re.search(r"\b(\d{1,2}\s+de\s+\w+(?:\s+de\s+\d{4})?|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b", part, re.I):
                        assigned_parts[date_key] = part
                        parts.remove(part)
                        break
            # 2. Buscar estado
            status_key = next((k for k in col_keys if "estado" in k or "status" in k or "avance" in k), None)
            if status_key:
                status_keywords = {"pendiente", "completado", "abierto", "cerrado", "en progreso", "ok", "en curso", "en proceso", "cancelado", "realizado", "hecho", "ejecutado"}
                for part in list(parts):
                    if part.lower() in status_keywords or any(kw in part.lower() for kw in ("no iniciado", "en desarrollo")):
                        assigned_parts[status_key] = part
                        parts.remove(part)
                        break
            # 3. Buscar responsable
            resp_key = next((k for k in col_keys if "responsable" in k or "encargado" in k or "nombre" in k or "autor" in k), None)
            if resp_key and len(parts) > 1:
                longest_part = max(parts, key=len)
                for part in list(parts):
                    if part != longest_part:
                        assigned_parts[resp_key] = part
                        parts.remove(part)
                        break
            # 4. El remanente va a la actividad principal
            if parts:
                res[primary_key] = parts[0]
            # Rellenar las clasificadas
            for k, v in assigned_parts.items():
                res[k] = v
        else:
            # ── Mapeo secuencial para participantes/tablas no estructuradas ──
            idx = 0
            for k in col_keys:
                if "firma" in k:
                    continue
                if idx < len(parts):
                    res[k] = parts[idx]
                    idx += 1
        return res

    # Caso 2: Texto plano sin comas
    # Extraer fecha si existe
    date_match = re.search(
        r"\b(\d{1,2}\s+de\s+\w+(?:\s+de\s+\d{4})?|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b",
        clean_item, re.I
    )
    extracted_date = ""
    if date_match:
        extracted_date = date_match.group(1)
        clean_item = clean_item.replace(date_match.group(0), "").strip()
        clean_item = re.sub(r"\b(para el|el|fecha|plazo)\s*$", "", clean_item, flags=re.I).strip()

    # Extraer responsable si hay dos puntos
    if ":" in clean_item:
        parts = clean_item.split(":", 1)
        res[primary_key] = parts[1].strip()
        resp_key = next((k for k in col_keys if "responsable" in k or "encargado" in k or "nombre" in k), None)
        if resp_key:
            res[resp_key] = parts[0].strip()
    else:
        if is_text_primary:
            res[primary_key] = clean_item
            # Buscar responsable por mayúsculas al inicio
            resp_match = re.search(r"^([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)+)\s+(?:se compromete|debe|realiza|a cargo de)\b", clean_item)
            if resp_match:
                resp_name = resp_match.group(1)
                resp_key = next((k for k in col_keys if "responsable" in k or "encargado" in k or "nombre" in k), None)
                if resp_key:
                    res[resp_key] = resp_name
                    res[primary_key] = clean_item.replace(resp_name, "", 1).strip()
        else:
            name_key = next((k for k in col_keys if "nombre" in k or "participante" in k or "asistente" in k), primary_key)
            res[name_key] = clean_item

    # Asignar la fecha si se extrajo
    date_key = next((k for k in col_keys if "fecha" in k or "plazo" in k or "limite" in k), None)
    if date_key and extracted_date:
        res[date_key] = extracted_date

    # Limpiar conectores en la actividad principal
    if is_text_primary:
        act_val = res[primary_key]
        act_val = re.sub(r"^(se compromete a|debe|a|para)\s+", "", act_val, flags=re.I).strip()
        res[primary_key] = act_val.capitalize()

    return res


def _fill_list_table(tbl, value, col_keys: list[str]):
    """Añade filas de datos a una tabla, después de la fila de headers.
    Usa los encabezados reales de la tabla para mapear columnas (ignora 'firma')."""
    if not value:
        return
    items = value if isinstance(value, list) else [value]

    # Headers reales de la primera fila (en minúsculas) para mapeo flexible
    real_headers = [c.text.strip().lower() for c in tbl.rows[0].cells] if tbl.rows else []

    template_row = tbl.rows[1] if len(tbl.rows) > 1 else tbl.rows[0]

    for item in items:
        new_row = copy.deepcopy(template_row._tr)
        tbl._tbl.append(new_row)
        row_obj = tbl.rows[-1]
        # Limpiar todo el texto de la fila clonada
        for cell in row_obj.cells:
            _set_cell_text(cell, "")

        if isinstance(item, str):
            item = _parse_string_item_to_cols(item, col_keys)

        if isinstance(item, dict):
            # Normalizar las claves del item (quitar acentos comunes, lowercase)
            item_norm = {k.lower().replace(" ", "_"): v for k, v in item.items()}

            for ci, header in enumerate(real_headers):
                if ci >= len(row_obj.cells):
                    break
                if "firma" in header:
                    continue  # columna de firma siempre vacía
                # Buscar el mejor match entre el header de la columna y las claves del item
                cell_val = ""
                for ik, iv in item_norm.items():
                    if ik in header or header in ik or _keys_similar(ik, header):
                        cell_val = str(iv or "")
                        break
                # Fallback: usar col_keys por posición
                if not cell_val and ci < len(col_keys):
                    cell_val = str(item_norm.get(col_keys[ci], ""))
                _set_cell_text(row_obj.cells[ci], cell_val)
        else:
            _set_cell_text(row_obj.cells[0], str(item))


def _fill_text_paragraphs(doc, flat: dict):
    """Busca párrafos en MAYÚSCULAS seguidos de vacío y rellena con el texto del acta."""
    paras = doc.paragraphs
    for idx, para in enumerate(paras):
        text = para.text.strip().rstrip(".")
        if not _is_section_heading(para):
            continue
        key = _to_key(text)
        value = flat.get(key)
        if value and idx + 1 < len(paras):
            next_para = paras[idx + 1]
            if not next_para.text.strip():
                next_para.text = str(value) if not isinstance(value, list) else "\n".join(str(v) for v in value)


def _set_cell_text(cell, text: str):
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.add_paragraph(text)


def _keys_similar(a: str, b: str) -> bool:
    """Compara claves con sinónimos frecuentes en actas corporativas."""
    aliases = {
        "entidad": ("organizacion", "empresa", "compania", "institution"),
        "nombre": ("participante", "persona", "asistente", "integrante"),
        "cargo": ("rol", "puesto", "posicion", "funcion"),
        "actividad": ("tarea", "compromiso", "accion", "item"),
        "responsable": ("encargado", "asignado", "responsability"),
        "estado": ("status", "avance", "progreso"),
        "descripcion": ("detalle", "contenido", "tema", "observacion"),
    }
    for canonical, syns in aliases.items():
        group = {canonical} | set(syns)
        if a in group and b in group:
            return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _to_key(text: str) -> str:
    return re.sub(r"\W+", "_", text.lower()).strip("_") or "campo"


def _infer_type(var: str, ctx: str) -> str:
    n, c = var.lower(), ctx.lower()
    if any(k in n for k in ("fecha", "date", "hora", "dia", "periodo", "plazo", "vencimiento")):
        return "date"
    if any(k in n for k in ("participante", "asistente", "integrante", "miembro", "equipo",
                             "firmante", "convocado", "nombre", "presente")):
        return "people"
    if any(k in n for k in ("lista", "items", "temas", "compromisos", "acuerdos",
                             "acciones", "tareas", "entregables", "agenda", "puntos")):
        return "list"
    if any(k in c for k in ("fecha", "date", "día", "hora")):
        return "date"
    if any(k in c for k in ("asistentes", "participantes", "presentes", "convocados")):
        return "people"
    return "text"


def _is_required(var: str, ctx: str) -> bool:
    keywords = ("fecha", "titulo", "tipo", "participante", "asistente",
                 "decision", "responsable", "acuerdo", "reunion", "objeto",
                 "lugar", "objetivo", "nombre")
    n, c = var.lower(), ctx.lower()
    return any(k in n or k in c for k in keywords)


def _build_hint(var: str, surrounding: str, ftype: str) -> str:
    name = var.replace("_", " ")
    if surrounding and surrounding.lower() != var.lower():
        return f'Buscar en la transcripción información sobre: "{surrounding}"'
    return {
        "date": f"Fecha o momento de {name}",
        "people": f"Personas o participantes de {name}",
        "list": f"Lista de elementos para {name}",
        "text": f"Información sobre {name}",
    }.get(ftype, name)
